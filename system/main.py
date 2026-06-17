from docx import Document
from datetime import datetime
import subprocess
import pandas as pd
import os

applicant = input("Enter the applicant's name: ")

df = pd.read_excel('items.xlsx')
df.columns = df.columns.str.strip()

cat_map = {
        'CATEGORY_A': ('☒','☐','☐'),
        'CATEGORY_B': ('☐','☒','☐'),
        'CATEGORY_C': ('☐','☒','☐'),
        'CATEGORY_D': ('☐','☐','☒')
    }

impact_map = {
        'HIGH IMPACT': ('☒','☐','☐'),
        'MEDIUM IMPACT': ('☐','☒','☐'),
        'LOW IMPACT': ('☐','☐','☒')
    }

priority_map = {
        'HIGH': ('☒','☐','☐'),
        'MEDIUM': ('☐','☒','☐'),
        'LOW': ('☐','☐','☒')
    }

for _, line in df.iterrows():
        
    desc = str(line['DESCRIPTION'])
    cat = str(line['CATEGORY'].upper())
    price = f"{float(line['PRICE']):.2f}".replace('.', ',')
    un = str(line['UNIT'])
    impact = str(line['IMPACT'].upper())
    priority = str(line['PRIORITY'].upper())
    now = datetime.now()

    doc = Document('template.docx')

    caa, cab, cad = cat_map.get(cat, ('☐','☐','☐'))
    hi, mi, li = impact_map.get(impact, ('☐','☐','☐'))
    h, m, l = priority_map.get(priority, ('☐','☐','☐'))

    elements = list(doc.paragraphs)

    for table in doc.tables:
       for row in table.rows:
        for cell in row.cells:
            elements.extend(cell.paragraphs)

    for paragraph in elements:

        text = paragraph.text

        text = text.replace('((applicant))', applicant)
        text = text.replace('((description))', desc)
        text = text.replace('((price))', price)
        text = text.replace('((un))', un)
        text = text.replace('((category))', cat)
        text = text.replace('((date))', now.strftime('%d/%m/%y'))

        text = text.replace('((caa))', caa)
        text = text.replace('((cab))', cab)
        text = text.replace('((cad))', cad)

        text = text.replace('((hi))', hi)
        text = text.replace('((mi))', mi)
        text = text.replace('((li))', li)

        text = text.replace('((h))', h)
        text = text.replace('((m))', m)
        text = text.replace('((l))', l)

        paragraph.text = text

    char = '<>:"/\\|?*'
    filename = desc.translate(
        str.maketrans(char, ' ' * len(char))
        )

    file = (filename + '.docx')
    doc.save(file)
    
    subprocess.run([
       'libreoffice',
       '--headless',
       '--convert-to', 'pdf',
       file
    ])

    os.remove(filename + '.docx')